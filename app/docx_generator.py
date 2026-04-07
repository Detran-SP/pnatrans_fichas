from __future__ import annotations

import re
from collections.abc import Callable
from datetime import datetime
from pathlib import Path
from urllib.parse import unquote

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.processing import sanitize_filename

PROJECT_ROOT = Path(__file__).resolve().parent.parent

# Paleta de cores (espelha o template Typst em main.qmd)
COLOR_PILAR = "004077"  # azul escuro - fundo do Pilar (H1)
COLOR_SECTION = "005CA8"  # azul médio - fundo das seções (H2)
COLOR_TABLE_HEAD = "A4CFED"  # azul claro - cabeçalho de tabela
COLOR_LINK = "0000FF"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT_NAME = "Open Sans"
FONT_SIZE = Pt(10)
HEADING_SIZE = Pt(14)


# ---------------------------------------------------------------------------
# Helpers de baixo nível (manipulam XML do python-docx)
# ---------------------------------------------------------------------------


def _set_paragraph_shading(paragraph, color_hex: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def _set_cell_shading(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _style_run(
    run, *, bold: bool = False, color: RGBColor | None = None, size: Pt | None = None
) -> None:
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Insere um hyperlink formatado (azul, sublinhado) no parágrafo."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")  # 10pt = 20 half-points
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLOR_LINK)
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Helpers de alto nível (adicionam blocos ao documento)
# ---------------------------------------------------------------------------


def _add_heading(doc: Document, text: str, bg_color: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _style_run(run, bold=True, color=WHITE, size=HEADING_SIZE)
    _set_paragraph_shading(p, bg_color)


def _add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run("" if pd.isna(text) else str(text))
    _style_run(run, bold=bold)


def _add_label_value(doc: Document, label: str, value) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(label)
    _style_run(label_run, bold=True)
    space_run = p.add_run(" ")
    _style_run(space_run)
    value_run = p.add_run("" if pd.isna(value) else str(value))
    _style_run(value_run)


def _add_links_list(doc: Document, links_string) -> None:
    """Renderiza uma string ';'-separada como lista numerada de hyperlinks."""
    if pd.isna(links_string) or str(links_string).strip() == "":
        _add_paragraph(doc, "Nenhum link fornecido")
        return

    links = [l.strip() for l in str(links_string).split(";") if l.strip()]
    if not links:
        _add_paragraph(doc, "Nenhum link fornecido")
        return

    for idx, url in enumerate(links, 1):
        match = re.search(r"[^/]+\.[a-zA-Z0-9]+$", url)
        filename = unquote(match.group()) if match else f"Link {idx}"
        p = doc.add_paragraph()
        prefix_run = p.add_run(f"{idx}. ")
        _style_run(prefix_run)
        _add_hyperlink(p, url, filename)


def _add_results_table(doc: Document, ano_ref, resultado, observacao) -> None:
    headers = ["Ano de Referência", "Valor do Resultado", "Observação"]
    values = [
        str(ano_ref),
        str(resultado),
        "" if pd.isna(observacao) else str(observacao),
    ]

    table = doc.add_table(rows=2, cols=3)
    table.autofit = True

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        _style_run(run, bold=True)
        _set_cell_shading(cell, COLOR_TABLE_HEAD)

    for i, value in enumerate(values):
        cell = table.rows[1].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(value)
        _style_run(run)


# ---------------------------------------------------------------------------
# Configuração do documento e construção por linha
# ---------------------------------------------------------------------------


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE


def _set_page_layout(doc: Document, logo_path: Path) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(1.27)

    if logo_path.exists():
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header_para.add_run()
        run.add_picture(str(logo_path), width=Cm(4.2))  # ~20% da largura da página


def _build_document(row: pd.Series, logo_path: Path) -> Document:
    doc = Document()
    _set_default_font(doc)
    _set_page_layout(doc, logo_path)

    tipo = row["tipo_de_produto_a_ser_enviado"]
    pilar = row["pilar"]
    acao_id = row["acao"]
    acao_name = row["nome_acao"] if pd.notna(row.get("nome_acao")) else ""
    acao_label = f"{acao_id}: {acao_name}" if acao_name else str(acao_id)
    prod_id = row["produto"]
    prod_name = row["nome_produto"] if pd.notna(row.get("nome_produto")) else ""
    prod_label = f"{prod_id}: {prod_name}" if prod_name else str(prod_id)
    prod_detran_name = row["titulo_do_produto_detran"]
    justificativa = row["descricao_e_justificativa"]
    indicadores = row["indicador"]
    ano_ref = int(float(row["ano_referencia"]))
    resultado = int(float(row["valor_resultado"]))
    observacao = row["observacoes"] if pd.notna(row.get("observacoes")) else ""
    area = row["area_responsavel"]
    links = row.get("links_comprovatorios")
    arquivos = row.get("arquivos_comprovatorios")

    # Pilar (H1)
    _add_heading(doc, str(pilar), COLOR_PILAR)

    # AÇÃO
    _add_heading(doc, "AÇÃO", COLOR_SECTION)
    _add_paragraph(doc, acao_label, bold=True)

    # PRODUTO
    if tipo == "Plano de Ação":
        _add_heading(doc, "PRODUTO", COLOR_SECTION)
        _add_paragraph(doc, prod_label, bold=True)
    else:  # "Produto Novo"
        _add_heading(doc, str(prod_id), COLOR_SECTION)

    _add_label_value(doc, "Título do Produto:", prod_detran_name)
    _add_label_value(doc, "Descrição e Justificativa:", justificativa)

    # INDICADORES
    _add_heading(doc, "INDICADORES", COLOR_SECTION)
    _add_paragraph(doc, indicadores)

    # RESULTADOS / METAS
    _add_heading(doc, "RESULTADOS / METAS", COLOR_SECTION)
    _add_results_table(doc, ano_ref, resultado, observacao)

    # ÁREA RESPONSÁVEL
    _add_heading(doc, "ÁREA RESPONSÁVEL PELA AÇÃO (Detran-SP)", COLOR_SECTION)
    _add_paragraph(doc, area)

    # ARQUIVOS COMPROBATÓRIOS
    _add_heading(doc, "ARQUIVOS COMPROBATÓRIOS DA EXECUÇÃO DA META", COLOR_SECTION)

    _add_paragraph(doc, "Links:", bold=True)
    _add_links_list(doc, links)

    _add_paragraph(doc, "Arquivos:", bold=True)
    _add_links_list(doc, arquivos)

    return doc


def generate_docx_documents(
    df: pd.DataFrame,
    output_dir: Path,
    on_progress: Callable | None = None,
) -> list[Path]:
    """Gera documentos .docx para cada linha do DataFrame."""
    output_dir.mkdir(parents=True, exist_ok=True)
    datetime_stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    logo_path = PROJECT_ROOT / "logo.png"
    generated_files: list[Path] = []

    for i in range(len(df)):
        row = df.iloc[i]
        produto_detran_clean = sanitize_filename(
            str(row.get("titulo_do_produto_detran", ""))
        )
        produto_clean = re.sub(r"\s+", "_", str(row.get("produto", "")))
        output_filename = (
            f"{row['acao']}_{produto_clean}_{produto_detran_clean}"
            f"_{datetime_stamp}.docx"
        )
        output_path = output_dir / output_filename

        doc = _build_document(row, logo_path)
        doc.save(str(output_path))
        generated_files.append(output_path)

        if on_progress:
            on_progress(i + 1, len(df))

    return generated_files
